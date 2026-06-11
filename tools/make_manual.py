#!/usr/bin/env python
"""Regenerate the LLM Duel manual end to end.

Captures fresh screenshots of the running app, then builds both the PDF and the
Word (.docx) manual from them.

Usage:
    python tools/make_manual.py                 # screenshots + PDF + DOCX
    python tools/make_manual.py --skip-capture  # rebuild PDF + DOCX from existing assets/

Requires:
    pip install streamlit playwright python-docx
    python -m playwright install chromium

Outputs (in the project root):
    assets/01_overview.png, 02_templates.png, 03_chat.png, 05_terminal.png
    LLM_Duel_Manual.pdf
    LLM_Duel_Manual.docx
"""
import argparse
import contextlib
import socket
import subprocess
import sys
import time
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
ASSETS = ROOT / "assets"
HTML = ROOT / "manual.html"
PDF = ROOT / "LLM_Duel_Manual.pdf"
DOCX = ROOT / "LLM_Duel_Manual.docx"
PORT = 8531
URL = f"http://localhost:{PORT}"

ACCENT = "5B4BDB"
INK = "1F2330"
MUTED = "6B7280"
CODE_FILL = "F4F5FA"
LINE = "D9DCE3"

# Hide Streamlit's dev chrome (top bar / toolbar / footer) for clean shots.
HIDE_CSS = """
[data-testid="stHeader"], [data-testid="stToolbar"],
[data-testid="stStatusWidget"], [data-testid="stDecoration"],
footer { display: none !important; }
"""

# Streamlit's chat layout auto-scrolls to the bottom; reset every scrollable
# region to the top so the page title is in frame.
RESET_SCROLL = """
() => {
  if (document.activeElement) { document.activeElement.blur(); }
  for (const el of document.querySelectorAll('*')) {
    const oy = getComputedStyle(el).overflowY;
    if (el.scrollHeight > el.clientHeight + 5 && (oy === 'auto' || oy === 'scroll')) {
      el.scrollTop = 0;
    }
  }
  window.scrollTo(0, 0);
}
"""

ANALYZER_CODE = (
    'response = f"Words: {len(prompt.split())} | Reversed: " '
    '+ " ".join(reversed(prompt.split()))'
)


# --------------------------------------------------------------------------- #
# 1. Screenshots                                                              #
# --------------------------------------------------------------------------- #
def _wait_for_port(port, timeout=60):
    end = time.time() + timeout
    while time.time() < end:
        with contextlib.closing(socket.socket()) as sock:
            sock.settimeout(1)
            if sock.connect_ex(("127.0.0.1", port)) == 0:
                return True
        time.sleep(0.5)
    return False


def capture_screens():
    from playwright.sync_api import sync_playwright

    ASSETS.mkdir(exist_ok=True)
    print("Starting Streamlit on port", PORT)
    proc = subprocess.Popen(
        [sys.executable, "-m", "streamlit", "run", str(ROOT / "app.py"),
         "--server.headless", "true", "--server.port", str(PORT),
         "--browser.gatherUsageStats", "false"],
        cwd=str(ROOT), stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )
    try:
        if not _wait_for_port(PORT):
            raise RuntimeError("Streamlit did not start; is it installed?")
        time.sleep(2)
        with sync_playwright() as pw:
            browser = pw.chromium.launch()
            page = browser.new_context(
                viewport={"width": 1500, "height": 1000}, device_scale_factor=2
            ).new_page()
            page.goto(URL, wait_until="domcontentloaded")
            page.get_by_placeholder("Message all panels").wait_for(timeout=45000)
            page.wait_for_timeout(2000)
            page.add_style_tag(content=HIDE_CSS)

            def reset_top():
                page.evaluate(RESET_SCROLL)
                time.sleep(0.4)

            def snap(name):
                page.add_style_tag(content=HIDE_CSS)
                time.sleep(0.2)
                page.screenshot(path=str(ASSETS / name))
                print("  saved", name)

            reset_top()
            snap("01_overview.png")

            try:
                reset_top()
                page.locator('[data-baseweb="select"]').first.click()
                page.wait_for_timeout(600)
                snap("02_templates.png")
                page.keyboard.press("Escape")
                page.wait_for_timeout(300)
            except Exception as exc:
                print("  templates shot skipped:", exc)

            try:
                areas = page.locator('[data-testid="stTextArea"] textarea')
                areas.nth(1).click()
                areas.nth(1).fill(ANALYZER_CODE)
                # The panel title boxes are the only inputs labelled "Display
                # name" (sidebar inputs carry their own labels), and the label
                # survives label_visibility="collapsed" as an aria-label.
                names = page.get_by_label("Display name")
                names.nth(0).fill("Echo model")
                names.nth(1).fill("Analyzer model")
            except Exception as exc:
                print("  panel config skipped:", exc)

            try:
                ci = page.get_by_placeholder("Message all panels")
                ci.click()
                ci.fill("In one sentence, what is a transformer in machine learning?")
                ci.press("Enter")
                page.get_by_text("You said:").first.wait_for(timeout=25000)
                page.wait_for_timeout(2000)
            except Exception as exc:
                print("  chat interaction issue:", exc)
            reset_top()
            snap("03_chat.png")

            try:
                page.get_by_text("Terminal", exact=False).first.click()
                page.wait_for_timeout(500)
                cmd = page.get_by_placeholder("pip install")
                cmd.click()
                cmd.fill("pip --version")
                page.keyboard.press("Tab")          # commit value to Streamlit
                page.wait_for_timeout(900)
                page.get_by_role("button", name="Run command").click()
                page.get_by_text("exit ", exact=False).first.wait_for(timeout=20000)
                page.wait_for_timeout(800)
                page.get_by_text("Terminal", exact=False).first.scroll_into_view_if_needed()
                page.wait_for_timeout(500)
                snap("05_terminal.png")
            except Exception as exc:
                print("  terminal shot skipped:", exc)

            browser.close()
    finally:
        proc.terminate()
        try:
            proc.wait(timeout=10)
        except Exception:
            proc.kill()
        print("Stopped Streamlit")


# --------------------------------------------------------------------------- #
# 2. PDF (Chromium print engine renders manual.html)                          #
# --------------------------------------------------------------------------- #
def build_pdf():
    from playwright.sync_api import sync_playwright

    footer = (
        '<div style="width:100%; font-size:8px; color:#9aa0aa; text-align:center; '
        'padding-top:2px;">LLM Duel — User Manual &nbsp;·&nbsp; '
        '<span class="pageNumber"></span> / <span class="totalPages"></span></div>'
    )
    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page()
        page.goto(HTML.as_uri(), wait_until="load")
        page.wait_for_timeout(800)
        page.pdf(
            path=str(PDF), format="A4", print_background=True,
            display_header_footer=True, header_template="<div></div>",
            footer_template=footer,
            margin={"top": "14mm", "bottom": "16mm", "left": "15mm", "right": "15mm"},
        )
        browser.close()
    print("wrote", PDF.name)


# --------------------------------------------------------------------------- #
# 3. DOCX (python-docx)                                                       #
# --------------------------------------------------------------------------- #
def build_docx():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    for section in doc.sections:        # US Letter is python-docx's default; set margins
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
    content_w = Inches(6.5)

    # python-docx's default settings.xml ships a bare <w:zoom/>; the schema wants
    # a percent. Set it so strict validators are happy.
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")

    def _shd(el_parent, fill):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        el_parent.append(shd)

    def box(paragraph, color, fill=None, size=6):
        pPr = paragraph._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement("w:" + edge)
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(size))
            e.set(qn("w:space"), "6")
            e.set(qn("w:color"), color)
            pbdr.append(e)
        pPr.append(pbdr)                # pBdr precedes shd in schema order
        if fill:
            _shd(pPr, fill)

    def rule(paragraph, color, size=10):
        pPr = paragraph._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "2")
        b.set(qn("w:color"), color)
        pbdr.append(b)
        pPr.append(pbdr)

    def mono(paragraph, text, size=10, color="333333"):
        run = paragraph.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        t = run._r.find(qn("w:t"))
        if t is not None:
            t.set(qn("xml:space"), "preserve")
        return run

    def rich(paragraph, segments):
        # segments: (text, kind) with kind in {None, "b", "code"}
        for text, kind in segments:
            if kind == "code":
                mono(paragraph, text)
            else:
                run = paragraph.add_run(text)
                if kind == "b":
                    run.bold = True
        return paragraph

    def code_block(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        for i, line in enumerate(text.split("\n")):
            if i:
                p.add_run().add_break()
            mono(p, line, size=9, color="2B2B2B")
        box(p, LINE, fill=CODE_FILL)
        return p

    def figure(name, caption):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ASSETS / name), width=content_w)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)
        cap.paragraph_format.space_after = Pt(12)

    def heading(text, level):
        h = doc.add_heading(text, level)
        return h

    # ---- Cover --------------------------------------------------------------
    title = doc.add_heading("LLM Duel", 0)
    sub = doc.add_paragraph()
    r = sub.add_run("User Manual — run one prompt against two LLMs, side by side.")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    meta = doc.add_paragraph()
    r = meta.add_run("Streamlit application  ·  June 2026")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    rule(meta, ACCENT, size=12)
    meta.paragraph_format.space_after = Pt(10)

    # ---- 1 ------------------------------------------------------------------
    heading("1  What it is", 1)
    rich(doc.add_paragraph(), [
        ("LLM Duel lets you hold the ", None), ("same conversation", "b"),
        (" with two (or more) language models at once and watch their answers ", None),
        ("stream in side by side", "b"),
        (". Each model is driven by a small block of Python you can edit in the app, "
         "so you can wire up ChatGPT, Claude, Gemini, a local model, or your own — "
         "without the app needing to know how any of them work.", None),
    ])
    figure("01_overview.png",
           "The main window: settings sidebar (left), two model panels each with an "
           "editable code box, and a shared message box at the bottom.")

    # ---- 2 ------------------------------------------------------------------
    heading("2  Install & run", 1)
    doc.add_paragraph("You need Python with Streamlit installed. From the project folder:")
    code_block("pip install -r requirements.txt\n"
               "copy .env.example .env      # then paste in your API keys (optional)")
    doc.add_paragraph("Then start it in any of these ways:")
    rich(doc.add_paragraph(style="List Bullet"),
         [("Double-click ", None), ("run.bat", "code"), (" — the quickest way.", None)])
    rich(doc.add_paragraph(style="List Bullet"),
         [("Or run ", None), ("streamlit run app.py", "code"), (" in a terminal.", None)])
    rich(doc.add_paragraph(), [
        ("Your browser opens at ", None), ("http://localhost:8501", "code"),
        (". It works immediately with the built-in ", None), ("Echo", "b"),
        (" template — no API keys required — so you can try the interface right away.", None),
    ])

    # ---- 3 ------------------------------------------------------------------
    heading("3  The interface", 1)
    heading("Sidebar (Settings)", 2)
    rich(doc.add_paragraph(style="List Bullet"),
         [("Panels", "b"), (" — compare 2, 3, or 4 models at once.", None)])
    rich(doc.add_paragraph(style="List Bullet"),
         [("Clear all chats", "b"), (" — wipe every conversation and start fresh.", None)])
    rich(doc.add_paragraph(style="List Bullet"),
         [("Save / Load / Download JSON", "b"),
          (" — store your panel names + code to ", None),
          ("llm_duel_config.json", "code"), (" so your setup survives a restart.", None)])
    heading("Each model panel", 2)
    rich(doc.add_paragraph(style="List Bullet"),
         [("Display name", "b"), (" — label the panel (e.g. “GPT-4o”, “Claude”).", None)])
    rich(doc.add_paragraph(style="List Bullet"),
         [("Code", "b"), (" — an editable Python box plus a template menu. It opens "
          "automatically and collapses once the conversation starts.", None)])
    rich(doc.add_paragraph(style="List Bullet"),
         [("Transcript", "b"), (" — the running conversation for that model.", None)])
    heading("Message box", 2)
    rich(doc.add_paragraph(),
         [("Type once at the bottom; your message goes to ", None), ("every", "b"),
          (" panel at the same time.", None)])

    # ---- 4 ------------------------------------------------------------------
    heading("4  How it works — the contract", 1)
    doc.add_paragraph(
        "The clever part is a tiny, fixed contract that every panel's code follows. "
        "That single convention is what lets one app drive any model:")
    rows = [
        ("Direction", "Name", "What it is", True),
        ("in", "prompt", "The latest user message (a string).", False),
        ("in", "messages", "The full conversation so far — a list of {\"role\", \"content\"} "
                           "dictionaries.", False),
        ("out", "response", "Either a plain string (shown at once), or a generator that "
                            "yields text chunks (streamed live).", False),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    widths = (Inches(1.1), Inches(1.3), Inches(4.1))
    for ri, (a, b, c, is_head) in enumerate(rows):
        cells = table.rows[ri].cells
        for ci, val in enumerate((a, b, c)):
            cells[ci].width = widths[ci]
            para = cells[ci].paragraphs[0]
            if is_head:
                run = para.add_run(val)
                run.bold = True
                _shd(cells[ci]._tc.get_or_add_tcPr(), "ECECF6")
            elif ci == 1:
                mono(para, val, size=10, color=ACCENT)
            else:
                para.add_run(val)
    doc.add_paragraph()
    rich(doc.add_paragraph(), [
        ("Pick a starting point from the template menu — ", None),
        ("OpenAI, Anthropic, Gemini, a local model (Ollama), or Echo", "b"),
        (" — then edit it. A real example (OpenAI, streaming, multi-turn):", None),
    ])
    code_block(
        'import os\n'
        'from openai import OpenAI\n\n'
        'client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))\n'
        'stream = client.chat.completions.create(\n'
        '    model="gpt-4o",\n'
        '    messages=messages,        # full history, so it remembers the conversation\n'
        '    stream=True,\n'
        ')\n\n'
        'def gen():\n'
        '    for event in stream:\n'
        '        if event.choices[0].delta.content:\n'
        '            yield event.choices[0].delta.content\n\n'
        'response = gen()              # a generator -> streamed token by token')
    figure("02_templates.png",
           "The template menu in each panel — a ready-made starting point for each provider.")

    # ---- 5 ------------------------------------------------------------------
    heading("5  Using it — a worked example", 1)
    for step in [
        "In each panel, pick a template from the Code menu and adjust the model name.",
        "Put your API keys in a .env file (see section 7).",
        "Type a question in the message box and press Enter.",
        "Watch both answers stream in, then compare. Ask a follow-up — each model "
        "remembers its own thread.",
    ]:
        doc.add_paragraph(step, style="List Number")
    rich(doc.add_paragraph(), [
        ("Below, two panels were given different code (a plain echo vs. a small text "
         "analyzer) so you can see how distinct each side's answer can be. The ", None),
        ("⏱", "b"), (" stamp under a reply shows how long it took.", None),
    ])
    figure("03_chat.png",
           "Same question, two panels, two answers — streamed live and shown side by side.")

    # ---- 6 ------------------------------------------------------------------
    heading("6  Installing model packages (Terminal)", 1)
    rich(doc.add_paragraph(), [
        ("Different models need different Python packages. The built-in ", None),
        ("Terminal", "b"),
        (" panel runs commands in the same interpreter the app uses, so anything you ", None),
        ("pip install", "code"), (" there is immediately importable by your panel code.", None),
    ])
    code_block("pip install openai anthropic google-generativeai")
    figure("05_terminal.png",
           "The Terminal panel running a command and showing its output and exit code.")

    # ---- 7 ------------------------------------------------------------------
    heading("7  Saving setups & API keys", 1)
    rich(doc.add_paragraph(), [
        ("Use ", None), ("Save", "b"), (" in the sidebar to write your panel names and code "
         "to ", None), ("llm_duel_config.json", "code"),
        ("; it loads automatically next time. ", None), ("Download JSON", "b"),
        (" gives you a portable copy.", None)])
    rich(doc.add_paragraph(), [
        ("API keys live in a ", None), (".env", "code"), (" file beside ", None),
        ("app.py", "code"), (" and are read by your panel code via ", None),
        ("os.environ", "code"), (":", None)])
    code_block("OPENAI_API_KEY=sk-...\nANTHROPIC_API_KEY=sk-ant-...\nGOOGLE_API_KEY=...")

    # ---- 8 ------------------------------------------------------------------
    heading("8  Security & troubleshooting", 1)
    callout = doc.add_paragraph()
    rich(callout, [
        ("Security: ", "b"),
        ("panels run with ", None), ("exec()", "code"),
        (" and the Terminal runs shell commands — full power, by design, so any SDK "
         "works. Only paste code you trust. This is built as a local, single-user tool.", None),
    ])
    box(callout, "F4B8AD", fill="FDECEA")
    callout.paragraph_format.space_after = Pt(8)
    for label, rest in [
        ("Browser didn't open:", " go to http://localhost:8501 manually."),
        ("ModuleNotFoundError in a panel:", " install the package from the Terminal panel, "
                                            "e.g. pip install openai."),
        ("Authentication error:", " check the matching key in .env, then restart the app."),
        ("“model not found”:", " edit the model=\"...\" line in that panel's code."),
        ("One panel is slow:", " that's just that model — panels stream independently and "
                              "in parallel."),
    ]:
        rich(doc.add_paragraph(style="List Bullet"), [(label, "b"), (rest, None)])

    # ---- Footer with page number -------------------------------------------
    foot = doc.sections[0].footer.paragraphs[0]
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run("LLM Duel — User Manual  ·  Page ")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    fld_run = foot.add_run()
    fld_run.font.size = Pt(8)
    fld_run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    fld_run._r.append(begin); fld_run._r.append(instr); fld_run._r.append(end)

    doc.save(str(DOCX))
    print("wrote", DOCX.name)


def main():
    parser = argparse.ArgumentParser(description="Rebuild the LLM Duel manual.")
    parser.add_argument("--skip-capture", action="store_true",
                        help="reuse existing screenshots in assets/")
    args = parser.parse_args()
    if not args.skip_capture:
        capture_screens()
    build_pdf()
    build_docx()
    print("Done.")


if __name__ == "__main__":
    main()
